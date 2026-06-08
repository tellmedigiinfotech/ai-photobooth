"""
Build a comprehensive .docx with all project documentation, every preset
prompt, and every image (backgrounds, outfit reference templates, brand).

Run from the project root:
    python3 scripts/build_docs.py
"""

import os
import io
import re
import tempfile
from pathlib import Path

from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent.parent
ASSETS = ROOT / "assets"
BACKGROUNDS = ASSETS / "backgrounds"
TEMPLATES = ASSETS / "templates"
BRAND = ASSETS / "brand"
OUT = ROOT / "AI_Photobooth_Documentation.docx"

TMP = Path(tempfile.mkdtemp(prefix="aiphotobooth_doc_"))


# ---------------------------------------------------------------------------
#  Preset data (mirrored from api/generate.js + app.js)
# ---------------------------------------------------------------------------

PRESETS = [
    {
        "id": 1,
        "name": "Khajuraho — Kandariya Mahadev",
        "description": "UNESCO-listed Chandela-era sandstone temples",
        "bg": "Jagdambi Temple , Kandariya Mahadev Temple.jpg",
        "setting": "ancient carved sandstone temple at Khajuraho",
        "male": "chandela-era hindu devotee in a traditional white cotton dhoti, a saffron uttariya (shoulder cloth) draped over one shoulder, and a rudraksha mala",
        "female": "chandela-era noblewoman in a traditional ivory silk saree with a narrow gold border, classical temple gold jewellery (jhumka earrings, a thin choker and bangles) and a small red bindi",
        "genders": None,
    },
    {
        "id": 2,
        "name": "Khajuraho — Lakshmana Temple",
        "description": "The finely carved 10th-century Chandela temple",
        "bg": "Lakshmana Temple IMG_9753-HDR.jpg",
        "setting": "10th-century Chandela sandstone temple at Khajuraho",
        "male": "chandela-era nobleman in a cream cotton dhoti and unstitched shoulder cloth, a simple gold armlet and a rudraksha mala around the neck",
        "female": "chandela-era noblewoman in a soft mustard silk saree with a gold border, traditional temple gold jewellery and a small red bindi",
        "genders": None,
    },
    {
        "id": 3,
        "name": "Orchha — Jahangir Mahal",
        "description": "17th-century Bundela palace, arched courtyards",
        "bg": "Jahangir Mahal 6 - Copy.jpg",
        "setting": "17th-century Bundela palace courtyard at Orchha",
        "male": "bundela rajput prince in a brocade cream angarkha, churidar pyjama, a colourful safa (rajput turban) and a kamarband (waist sash), with subtle gold jewellery",
        "female": "bundela rajput princess in a deep-red silk lehenga-choli with zari embroidery, a sheer dupatta draped over the head, traditional rajput gold jewellery (maang tikka, jhumkas, choker)",
        "genders": None,
    },
    {
        "id": 4,
        "name": "Orchha — Jahangir Gate",
        "description": "Monumental Bundela-Mughal archway",
        "bg": "jahangir gate orchha.jpg",
        "setting": "monumental Mughal-era Bundela gateway at Orchha",
        "male": "bundela rajput warrior in a saffron brocade angarkha, churidar pyjama, a colourful safa turban, a kamarband and a scabbarded sword at the waist",
        "female": "bundela rajput princess in a maroon silk lehenga with gold zari work, a sheer dupatta draped over the head and traditional rajput gold jewellery",
        "genders": None,
    },
    {
        "id": 6,
        "name": "Maheshwar — Chhatri by the River",
        "description": "Holkar cenotaphs above the Narmada ghats",
        "bg": "Chattei River view (7).jpg",
        "setting": "18th-century Maratha riverside chhatri on the Narmada at Maheshwar",
        "male": "ahilyabai-era maratha nobleman in a crisp white cotton dhoti and bandgala-style kurta with a bright red pheta (maratha turban) and a simple shawl over one shoulder",
        "female": "ahilyabai-era maharashtrian lady in a traditional burgundy nauvari saree (9-yard drape), a gold nath (curved nose ring), thushi (short gold choker), green glass bangles, a pearl necklace and a small red decorative bindi",
        "genders": None,
    },
    {
        "id": 9,
        "name": "Krishnabai Holkar Chhatri",
        "description": "The queen's cenotaph above the Narmada, Maheshwar",
        "bg": "Krishnabai holkar chhatri .jpg",
        "setting": "18th-century Holkar royal chhatri at Maheshwar",
        "male": "holkar-era maratha sardar in a cream cotton dhoti-kurta with a red pheta turban and a shawl draped over one shoulder",
        "female": "holkar-era maharashtrian queen in a royal-blue nauvari saree with gold border, a gold nath, thushi, pearl necklace and ornate traditional jewellery, styled after queen ahilyabai holkar",
        "genders": None,
    },
    {
        "id": 10,
        "name": "Indore — Rajwada Palace",
        "description": "The seven-storey Holkar palace of Indore",
        "bg": "Rajwada Indore.jpg",
        "setting": "18th-century seven-storey Holkar palace in Indore",
        "male": "holkar-era maharaja in a cream brocade angarkha, churidar pyjama, a jewelled red pheta turban with a sarpech ornament, a kamarband and a pearl necklace",
        "female": "holkar-era maharani in a peacock-green paithani saree with a heavy gold zari border, a large gold nath, ornate thushi, a multi-strand pearl necklace, maang tikka and traditional regal jewellery",
        "genders": None,
    },
    {
        "id": 11,
        "name": "Indore — Rajwada Courtyard",
        "description": "Inside the Holkar royal seat",
        "bg": "RajWada 15.jpg",
        "setting": "inner courtyard of the 18th-century Holkar palace in Indore",
        "male": "holkar-era maratha nobleman in a cream cotton dhoti-kurta with a red pheta turban and a simple shawl",
        "female": "holkar-era royal lady in a teal nauvari saree with a gold border, a gold nath, thushi, pearl necklace and traditional maharashtrian jewellery",
        "genders": None,
    },
    {
        "id": 12,
        "name": "Kheoni Sanctuary — Wilds of MP",
        "description": "Central Indian teak and sal forest",
        "bg": "kheoni wildlife sanctuary .jpg",
        "setting": "central Indian teak and sal forest at Kheoni Wildlife Sanctuary",
        "male": "modern wildlife safari explorer in a clean khaki short-sleeve shirt with a chest pocket, light beige cargo trousers, a wide-brim canvas safari hat and binoculars hanging around the neck",
        "female": "modern wildlife safari explorer in a clean khaki short-sleeve shirt with a chest pocket, light beige cargo trousers, a wide-brim canvas safari hat and binoculars hanging around the neck",
        "genders": ["male"],
    },
    {
        "id": 13,
        "name": "Kheoni Sanctuary — Forest Trail",
        "description": "Quiet woodland of teak, sal and bamboo",
        "bg": "kheoni wildlife sanctuary 1.jpg",
        "setting": "forest trail through teak and bamboo at Kheoni Wildlife Sanctuary",
        "male": "modern wildlife safari explorer in a sand-beige short-sleeve shirt with a chest pocket, khaki cargo trousers, a wide-brim canvas safari hat and binoculars hanging around the neck",
        "female": "modern wildlife safari explorer in a sand-beige short-sleeve shirt with a chest pocket, khaki cargo trousers, a wide-brim canvas safari hat and binoculars hanging around the neck",
        "genders": ["male"],
    },
    {
        "id": 14,
        "name": "Goa — Cabo de Rama Beach",
        "description": "Palm-fringed Goan coast at golden-hour sunset",
        "bg": "Cabo de Rama Beach_DSC9670.jpg",
        "setting": "wide curving Goan beach at Cabo de Rama during a soft pink-and-orange sunset, with the calm Arabian Sea, golden sand, dark coastal rocks and tall coconut palms swaying along the shoreline",
        "male": "easy-going beach-day visitor in a soft cream linen full-sleeve shirt with the sleeves loosely rolled up to the forearm, fully buttoned, comfortable beige cotton beach trousers, and a pair of simple rubber flip-flop beach slippers — modest, fully clothed, no shorts and no swimwear",
        "female": "easy-going beach-day visitor in an ankle-length flowy white-and-pastel-floral cotton maxi dress with three-quarter sleeves and a high modest neckline, a wide-brim natural straw sun hat, and a pair of simple rubber flip-flop beach slippers — modest, fully covered, ankle-length skirt, no swimwear, no exposed shoulders or midriff",
        "genders": None,
    },
    {
        "id": 15,
        "name": "Goa — Cola Beach",
        "description": "Rocky Goan shoreline framed by forested hills",
        "bg": "Cola Beach_DSC9401.jpg",
        "setting": "secluded Goan cove at Cola Beach with shallow turquoise water lapping over weathered black-and-rust coastal rocks, a small strip of golden sand and a lush forested hillside rising behind",
        "male": "easy-going beach-day visitor in a soft sky-blue linen full-sleeve shirt with the sleeves loosely rolled up, fully buttoned, light sand-coloured cotton beach trousers, and a pair of simple rubber flip-flop beach slippers — modest, fully clothed, no shorts and no swimwear",
        "female": "easy-going beach-day visitor in an ankle-length flowy mint-green cotton maxi dress with three-quarter sleeves and a high modest neckline, a light sheer cotton scarf draped over the shoulders, and a pair of simple rubber flip-flop beach slippers — modest, fully covered, ankle-length skirt, no swimwear, no exposed shoulders or midriff",
        "genders": None,
    },
    {
        "id": 16,
        "name": "Goa — Salim Ali Bird Sanctuary",
        "description": "Chorão Island's serene mangrove wetland",
        "bg": "Dr. Salim Ali Bird Sanctuary_DSC8234.jpg",
        "setting": "peaceful mangrove forest and tidal creek at the Dr. Salim Ali Bird Sanctuary on Chorão Island, Goa, with arching mangrove branches, a leafy green canopy and still water reflecting the trees",
        "male": "modern Goan-day birdwatcher in a clean light-olive long-sleeve cotton shirt with a chest pocket, comfortable beige cotton trousers, lightweight canvas walking shoes and a pair of binoculars hanging around the neck — neat, modest, fully clothed smart-casual",
        "female": "modern Goan-day birdwatcher in a clean ivory long-sleeve cotton shirt with a chest pocket, comfortable beige cotton trousers, lightweight canvas walking shoes and a pair of binoculars hanging around the neck — neat, modest, fully clothed smart-casual",
        "genders": None,
    },
    {
        "id": 17,
        "name": "Gulmarg — St. Mary's Church",
        "description": "Heritage wooden church in a Kashmir alpine meadow",
        "bg": "Gulmarg landscapes .jpg",
        "setting": "historic wooden Maharani St. Mary's Church set in a sunlit Gulmarg alpine meadow in Kashmir, with daisies and a tall conifer beside it",
        "male": "very well-dressed gentleman in a tailored charcoal-grey three-piece wool suit, a crisp white shirt, a deep-burgundy silk tie, a neatly folded white pocket square and polished black oxford shoes — formal, refined and fully covered",
        "female": "very well-dressed lady in an elegant tea-length midi dress in deep emerald with three-quarter sleeves and a modest high neckline, layered with a fitted camel wool overcoat, simple pearl earrings and refined low-heeled shoes — formal, graceful and fully covered",
        "genders": None,
    },
    {
        "id": 19,
        "name": "Gulmarg — Daisy Field",
        "description": "Open pasture of daisies under Kashmir skies",
        "bg": "Gulmarg landscapes 3.jpg",
        "setting": "open Gulmarg alpine pasture carpeted with white daisies under a clear blue sky, with a single tree on the horizon, Kashmir",
        "male": "very well-dressed gentleman in a smart light-grey wool suit, a soft pastel-blue shirt, a navy silk tie and polished oxford shoes — formal, refined and fully covered",
        "female": "very well-dressed lady in an elegant pastel-blue midi dress with three-quarter sleeves and a modest high neckline, layered with a tailored ivory overcoat, simple pearl earrings and refined low-heeled shoes — formal, graceful and fully covered",
        "genders": None,
    },
]


PROMPT_TEMPLATE = """ABSOLUTE RULE — symbols of marital status. do NOT add any of the following to the person in the output UNLESS that exact symbol is clearly, visibly present on the person in the first reference photo:
  • sindoor (red or orange vermilion powder in the hair parting)
  • mangalsutra (black-bead and gold marriage necklace)
  • kumkum at the parting or forehead
  • bridal makeup, heavy nath / nose ring, or any other suhaag / saubhagya symbol
if the reference photo does not show these, the output must not show them — regardless of what the outfit description below suggests, regardless of the cultural setting, and regardless of what is "traditional" for the era or region. inventing these on a person who does not wear them (a child, an unmarried woman, a person of any age or background who simply does not use them) is a critical failure that causes religious and cultural offence. when in doubt, OMIT.

face-preservation rule: the face in the first reference photo must be preserved exactly in the output. same eyes, same nose, same mouth, same jawline, same skin tone, same age, same expression — every facial detail must be identical to the reference. do not beautify, smooth, slim, stylise, de-age, lighten or reshape the face in any way. if the face does not match the reference exactly, the image is wrong.

age-appropriateness rule: study the apparent age in the reference photo and adapt the outfit accordingly. if the reference shows a child, young girl or teenager, keep the attire age-appropriate — simpler, lighter jewellery; no nath / nose ring; no heavy bridal ornaments; no adult makeup; absolutely no sindoor, mangalsutra or kumkum. treat the outfit description below as a stylistic direction, not a literal checklist — drop any element that is not appropriate for the person's apparent age. a small plain decorative bindi is acceptable only if the person clearly appears to be an adult woman; otherwise omit it.

now create an image of this person standing in this {SETTING}, dressed like a {OUTFIT}. please adjust the lights and shadows so the person blends naturally into the scene. the image should look super realistic and natural."""


FACESWAP_PROMPT_FALLBACK = "a photorealistic portrait photograph of a person at a heritage location, natural lighting, sharp focus, authentic cultural attire, medium shot from the waist up"
FACESWAP_NEGATIVE = "cartoon, anime, painting, illustration, 3d render, plastic skin, over-smoothed, blurry face, distorted face, extra limbs, deformed, low quality, watermark, text, signature"


# ---------------------------------------------------------------------------
#  Helpers
# ---------------------------------------------------------------------------

def resize_for_doc(src: Path, max_w: int = 1400, quality: int = 82) -> Path:
    """Resize/recompress the image to a temp jpg and return the temp path."""
    out = TMP / f"{src.stem}__{max_w}.jpg"
    if out.exists():
        return out
    with Image.open(src) as im:
        im = im.convert("RGB")
        if im.width > max_w:
            ratio = max_w / im.width
            im = im.resize((max_w, int(im.height * ratio)), Image.LANCZOS)
        im.save(out, "JPEG", quality=quality, optimize=True)
    return out


def set_cell_shading(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Inter"
    return h


def add_paragraph(doc, text, *, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Inter"
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    run.font.name = "Menlo"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("1f2937")
    return p


def add_image_centered(doc, path: Path, width_in: float = 5.5, caption: str | None = None):
    if not path.exists():
        add_paragraph(doc, f"[Missing image: {path.name}]", italic=True, color="9f1239")
        return
    small = resize_for_doc(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(small), width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor.from_string("64748b")


def add_two_image_row(doc, left: Path, right: Path, left_caption: str, right_caption: str, width_in: float = 2.9):
    t = doc.add_table(rows=2, cols=2)
    t.autofit = True
    for col_idx, (path, caption) in enumerate([(left, left_caption), (right, right_caption)]):
        cell = t.cell(0, col_idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path and path.exists():
            small = resize_for_doc(path, max_w=900)
            run = para.add_run()
            run.add_picture(str(small), width=Inches(width_in))
        else:
            r = para.add_run("[no reference image]")
            r.italic = True
            r.font.color.rgb = RGBColor.from_string("94a3b8")

        cap_cell = t.cell(1, col_idx)
        cap_para = cap_cell.paragraphs[0]
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap_para.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor.from_string("64748b")


def add_kv_table(doc, rows):
    """rows = [(label, value), ...]"""
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = True
    for i, (k, v) in enumerate(rows):
        kc = t.cell(i, 0)
        vc = t.cell(i, 1)
        set_cell_shading(kc, "f1f5f9")
        kc.width = Inches(1.6)
        kp = kc.paragraphs[0]
        kr = kp.add_run(k)
        kr.bold = True
        kr.font.size = Pt(10)
        vp = vc.paragraphs[0]
        vr = vp.add_run(v)
        vr.font.size = Pt(10)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "cbd5e1")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
#  Build
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Inter"
    style.font.size = Pt(11)

    # -- Cover -----------------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("AI Photobooth")
    tr.font.size = Pt(40)
    tr.bold = True
    tr.font.color.rgb = RGBColor.from_string("a63e1e")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Madhya Pradesh Heritage · Project Documentation")
    sr.font.size = Pt(16)
    sr.italic = True
    sr.font.color.rgb = RGBColor.from_string("44403c")

    doc.add_paragraph()

    logo_path = BRAND / "aakhon-dekha-logo.png"
    if logo_path.exists():
        add_image_centered(doc, logo_path, width_in=3.5, caption="Brand mark — Aakhon Dekha")

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        "A web-based AI photobooth: capture a portrait, pick a heritage destination, "
        "and Gemini 3 generates a photoreal image of you placed at that location in "
        "period-appropriate attire. Generated photos are stamped with the brand logo "
        "and Aakhon Dekha experience locations."
    )
    mr.font.size = Pt(11)
    mr.italic = True
    mr.font.color.rgb = RGBColor.from_string("44403c")

    doc.add_page_break()

    # -- Overview --------------------------------------------------------
    add_heading(doc, "1. Project overview", 1)

    add_paragraph(doc,
        "AI Photobooth is a three-step web wizard — Capture → Choose → Reveal — that takes "
        "a webcam selfie, lets the user pick a heritage destination in Madhya Pradesh (and "
        "now Goa & Kashmir), and uses Google's Gemini 3 Pro Image model to compose the user "
        "into that location in setting-appropriate clothing. The final image is brand-stamped "
        "client-side with a logo plate and the partner's experience-location pill."
    )

    add_heading(doc, "Hosting", 3)
    add_paragraph(doc,
        "Production runs on Vercel (vercel.json declares two serverless functions in /api, "
        "each with a 60-second maxDuration). Earlier firebase.json / functions/ scaffolding "
        "remains in the repo from a prior hosting experiment but is not the live target."
    )

    add_heading(doc, "Core features", 3)
    for bullet in [
        "Three-step wizard UI with stepper progress (Capture → Choose destination → Reveal).",
        "Live webcam capture with multi-camera selection and mirrored-preview output.",
        "15 active heritage destinations across MP (Khajuraho, Orchha, Maheshwar, Indore, "
        "Kheoni), Goa (Cabo de Rama, Cola Beach, Salim Ali Sanctuary) and Kashmir (Gulmarg).",
        "Gender toggle (Male / Female) — switches outfit prompts and filters gender-restricted "
        "destinations (Kheoni Sanctuary is male-only on the current client).",
        "Server-side prompt construction with absolute safety rules for sindoor / mangalsutra, "
        "face preservation, and age-appropriate attire.",
        "Gemini 3 Pro Image generation at 2K, 3:4 aspect ratio.",
        "Client-side brand overlay (logo plate top-left + experience-locations pill "
        "bottom-center) composited onto the generated image at full resolution.",
        "Download, WhatsApp share (Web Share API on mobile, link + download fallback on "
        "desktop) and one-click print.",
        "An optional Replicate / InstantID face-swap endpoint (api/faceswap.js) that can be "
        "wired in to lock identity even harder.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(bullet)
        r.font.size = Pt(11)

    doc.add_page_break()

    # -- Architecture ----------------------------------------------------
    add_heading(doc, "2. Architecture & tech stack", 1)

    add_kv_table(doc, [
        ("Frontend", "Vanilla HTML / CSS / JS (no framework). Index → app.js wires the wizard."),
        ("Backend", "Vercel serverless functions in /api (Node 18+). Express server.js remains for local dev."),
        ("AI — image gen", "Google Gemini 3 Pro Image preview (model: gemini-3-pro-image-preview) via @google/genai."),
        ("AI — face swap (opt.)", "Replicate · zsxkib/instant-id (SDXL + IdentityNet + IP-Adapter)."),
        ("Image upload", "Multer (memory storage, 30 MB cap)."),
        ("Branding", "Composed client-side on a <canvas> — deterministic, no model overhead."),
        ("Fonts", "Fraunces (display) + Inter (UI) via Google Fonts."),
        ("Hosting", "Vercel (production). Firebase Hosting scaffolding present but unused."),
    ])

    add_heading(doc, "File layout", 3)
    add_code(doc, """AI photobooth/
├── index.html                    Main wizard markup (3 steps)
├── styles.css                    Heritage theme + responsive
├── app.js                        Frontend state machine + brand overlay
├── server.js                     Local Express dev server (legacy Nano-Banana flow)
├── package.json                  Deps: @google/genai, multer, replicate, express, cors
├── vercel.json                   Function maxDuration: 60s
├── firebase.json / .firebaserc   Unused hosting config (prior experiment)
├── api/
│   ├── generate.js               Primary endpoint — Gemini 3 Pro Image
│   ├── faceswap.js               Optional InstantID identity-lock endpoint
│   └── health.js                 Reports whether GEMINI / REPLICATE keys are set
├── assets/
│   ├── backgrounds/              Heritage location photos (18 files)
│   ├── templates/                Outfit reference photos, male/female pairs
│   └── brand/aakhon-dekha-logo.png
├── scripts/build_docs.py         This documentation generator
└── test_*.{js,py}                Standalone API smoke tests""")

    add_heading(doc, "Data flow — generation", 3)
    add_code(doc, """ user
  │
  │ 1. capture
  ▼
[ webcam ] ─► canvas.toBlob() ─► state.capturedBlob
  │
  │ 2. pick gender + destination
  ▼
[ POST /api/generate ]   multipart: userImage, presetId, gender
  │
  │ 3. server-side
  ▼
api/generate.js
  ├─ multer parses upload (Vercel-quirk shim for pre-buffered body)
  ├─ resolves PRESETS[presetId] → setting + male/female outfit
  ├─ buildPrompt() — safety rules + setting + outfit
  ├─ fetches background JPG from /assets/backgrounds via x-forwarded-host
  ├─ ai.models.generateContent({
  │     model: 'gemini-3-pro-image-preview',
  │     contents: [userPhoto, backgroundPhoto, prompt],
  │     responseModalities: ['Image'],
  │     imageConfig: { aspectRatio: '3:4', imageSize: '2K' }
  │   })
  └─ returns { generatedImage: <base64>, mimeType }
  │
  │ 4. client
  ▼
brandifyImage(dataUrl)
  ├─ draws photo to canvas
  ├─ overlays logo plate (top-left)
  └─ overlays locations pill (bottom-center)
  │
  ▼
[ download | WhatsApp | print ]""")

    doc.add_page_break()

    # -- User flow -------------------------------------------------------
    add_heading(doc, "3. User flow", 1)

    add_heading(doc, "Step 1 — Capture", 2)
    add_paragraph(doc,
        "User taps 'Start camera' to enable webcam (1080p ideal, user-facing). The viewfinder "
        "shows a head-and-shoulders guide. Multiple cameras populate a picker once permission "
        "is granted. On capture, the frame is drawn mirrored to match the live preview, then "
        "exported as JPEG at 0.95 quality. The user can retake or proceed."
    )

    add_heading(doc, "Step 2 — Choose destination", 2)
    add_paragraph(doc,
        "User first selects a gender (Male / Female), which sets the outfit lane and hides "
        "gender-restricted presets (Kheoni Sanctuary is male-only currently). The destinations "
        "grid renders preset cards with background thumbnails. Selecting a card enables the "
        "Generate button in the sticky action bar."
    )

    add_heading(doc, "Step 3 — Reveal", 2)
    add_paragraph(doc,
        "On submit, the client compresses the photo to ≤2048px JPEG (q=0.95) and POSTs "
        "multipart to /api/generate. A loading screen rotates through five hint phrases every "
        "4.5 s. When the base64 image returns it is fed through brandifyImage(), then "
        "displayed in a result frame with Start over / Download / WhatsApp / Print actions."
    )

    doc.add_page_break()

    # -- Brand overlay ---------------------------------------------------
    add_heading(doc, "4. Brand overlay", 1)

    add_paragraph(doc,
        "The brand mark and experience-locations strip are composited onto the generated "
        "image entirely on the client (app.js → brandifyImage). This avoids spending model "
        "tokens on text rendering, keeps output deterministic, and works at full 2K."
    )

    add_image_centered(doc, logo_path, width_in=3.0, caption="assets/brand/aakhon-dekha-logo.png")

    add_heading(doc, "Composition details", 3)
    add_kv_table(doc, [
        ("Logo plate", "Cream rounded rectangle, ~26% of image width, top-left at 3.5% / 2.8% inset. 7% inner padding. Soft drop shadow."),
        ("Logo source", "assets/brand/aakhon-dekha-logo.png — preloaded once via loadImage()."),
        ("Locations pill", "Dark translucent pill, 62% opacity black, bottom-center at 4% inset."),
        ("Locations text", "Bhopal · Orchha · Maheshwar · Boat Club (hardcoded in BRAND_LOCATIONS)."),
        ("Typography", "500-weight Inter, font size = max(18, ~2.2% of image height)."),
        ("Fallback", "If overlay throws, the raw Gemini image is shown unchanged (warning logged)."),
    ])

    doc.add_page_break()

    # -- Master prompt ---------------------------------------------------
    add_heading(doc, "5. The master prompt", 1)

    add_paragraph(doc,
        "Every generation runs the same prompt template (api/generate.js → buildPrompt). "
        "Two slots are filled per-request: {SETTING} and {OUTFIT}. The safety rules sit at "
        "the top of the prompt by design — once Gemini drifted onto outfit detail, marital "
        "symbols (sindoor, mangalsutra) crept in regardless of the source photo. Front-loading "
        "the rule fixed it."
    )

    add_heading(doc, "Template (verbatim)", 3)
    add_code(doc, PROMPT_TEMPLATE.replace("{SETTING}", "{SETTING}").replace("{OUTFIT}", "{OUTFIT}"))

    add_heading(doc, "Model call", 3)
    add_code(doc, """ai.models.generateContent({
  model: 'gemini-3-pro-image-preview',
  contents: [
    { inlineData: { mimeType: 'image/jpeg', data: <user photo base64> } },
    { inlineData: { mimeType: 'image/jpeg', data: <heritage background base64> } },
    { text: <built prompt> },
  ],
  config: {
    responseModalities: ['Image'],
    imageConfig: {
      aspectRatio: '3:4',
      imageSize: '2K',
    },
  },
});""")

    doc.add_page_break()

    # -- Per-preset gallery + prompts -----------------------------------
    add_heading(doc, "6. Destinations — backgrounds, outfits, prompts", 1)

    add_paragraph(doc,
        "Each preset combines a heritage background photo (sent to Gemini as the second "
        "reference image), a one-line setting description (the {SETTING} slot in the prompt), "
        "and a male / female outfit description (the {OUTFIT} slot, picked by the user's "
        "gender toggle). Where available, the original outfit reference photo shown to the "
        "designer is also included for visual reference — these are NOT sent to the model."
    )

    for p in PRESETS:
        add_horizontal_rule(doc)
        title = f"Preset {p['id']} · {p['name']}"
        add_heading(doc, title, 2)
        add_paragraph(doc, p["description"], italic=True, color="64748b")

        gend_note = "Male & Female" if not p["genders"] else " · ".join(p["genders"]).title() + " only"
        add_kv_table(doc, [
            ("Background file", p["bg"]),
            ("Setting (prompt slot)", p["setting"]),
            ("Gender availability", gend_note),
        ])

        bg_path = BACKGROUNDS / p["bg"]
        add_image_centered(doc, bg_path, width_in=6.0, caption=f"Heritage background — {p['bg']}")

        # Outfit reference templates (if any)
        male_t = TEMPLATES / f"{p['id']}-male.jpg"
        female_t = TEMPLATES / f"{p['id']}-female.jpg"
        if male_t.exists() or female_t.exists():
            add_heading(doc, "Outfit reference photos (designer-only, not sent to the model)", 3)
            add_two_image_row(
                doc,
                male_t if male_t.exists() else None,
                female_t if female_t.exists() else None,
                f"Male reference · {male_t.name}",
                f"Female reference · {female_t.name}",
            )
        else:
            add_paragraph(doc, "No outfit reference template photos for this preset.", italic=True, color="94a3b8")

        # Prompts
        add_heading(doc, "Male prompt — outfit slot", 3)
        add_code(doc, p["male"])

        add_heading(doc, "Female prompt — outfit slot", 3)
        add_code(doc, p["female"])

        add_heading(doc, "Resolved full prompt — Male", 3)
        add_code(doc, PROMPT_TEMPLATE.replace("{SETTING}", p["setting"]).replace("{OUTFIT}", p["male"]))

        add_heading(doc, "Resolved full prompt — Female", 3)
        add_code(doc, PROMPT_TEMPLATE.replace("{SETTING}", p["setting"]).replace("{OUTFIT}", p["female"]))

    doc.add_page_break()

    # -- Inactive / unreferenced assets ---------------------------------
    add_heading(doc, "7. Inactive & unreferenced assets", 1)

    used_bgs = {p["bg"] for p in PRESETS}
    all_bgs = sorted([f.name for f in BACKGROUNDS.iterdir() if f.is_file() and not f.name.startswith(".")])
    unused_bgs = [b for b in all_bgs if b not in used_bgs]

    add_heading(doc, "Backgrounds present but not wired into the current preset list", 3)
    if not unused_bgs:
        add_paragraph(doc, "None — every background in assets/backgrounds is in use.", italic=True)
    else:
        add_paragraph(doc, "These files sit in assets/backgrounds/ but no entry in PRESETS references them. They may be drafts, retired destinations, or staged for a future release.")
        for b in unused_bgs:
            bp = BACKGROUNDS / b
            add_paragraph(doc, b, bold=True)
            add_image_centered(doc, bp, width_in=5.0, caption=b)

    add_heading(doc, "Outfit reference templates without a matching active preset", 3)
    active_ids = {p["id"] for p in PRESETS}
    orphan_templates = []
    for f in sorted(TEMPLATES.iterdir()):
        if not f.is_file():
            continue
        m = re.match(r"(\d+)-(male|female)\.jpg$", f.name)
        if not m:
            continue
        tid = int(m.group(1))
        if tid not in active_ids:
            orphan_templates.append(f)
    if not orphan_templates:
        add_paragraph(doc, "None — every template in assets/templates maps to an active preset.", italic=True)
    else:
        add_paragraph(doc, "These outfit references exist in assets/templates/ but their preset id is not in the active PRESETS map (likely retired destinations such as Royal Palace, Taj Mahal, Mysore, Hawa Mahal listed in the earlier README).")
        # Pair them up by id
        by_id = {}
        for f in orphan_templates:
            tid = int(re.match(r"(\d+)-", f.name).group(1))
            by_id.setdefault(tid, {})[("male" if "-male" in f.name else "female")] = f
        for tid in sorted(by_id):
            add_paragraph(doc, f"Orphan id {tid}", bold=True)
            add_two_image_row(
                doc,
                by_id[tid].get("male"),
                by_id[tid].get("female"),
                f"{tid}-male.jpg" if by_id[tid].get("male") else "(missing)",
                f"{tid}-female.jpg" if by_id[tid].get("female") else "(missing)",
            )

    doc.add_page_break()

    # -- All backgrounds gallery (used) ---------------------------------
    add_heading(doc, "8. All active heritage backgrounds — gallery", 1)
    add_paragraph(doc, "Every background image referenced by the active preset list, in preset order.")
    for p in PRESETS:
        bp = BACKGROUNDS / p["bg"]
        add_image_centered(doc, bp, width_in=6.0, caption=f"{p['id']} · {p['name']} — {p['bg']}")

    doc.add_page_break()

    # -- All outfit reference templates gallery -------------------------
    add_heading(doc, "9. All outfit reference photos — gallery", 1)
    add_paragraph(doc,
        "Designer reference photos used to author the per-preset outfit prompts. They live "
        "in assets/templates/ as {presetId}-{male|female}.jpg. They are NOT sent to Gemini — "
        "only the user's webcam capture and the heritage background are."
    )
    seen_ids = set()
    for p in PRESETS:
        male_t = TEMPLATES / f"{p['id']}-male.jpg"
        female_t = TEMPLATES / f"{p['id']}-female.jpg"
        if male_t.exists() or female_t.exists():
            add_heading(doc, f"Preset {p['id']} — {p['name']}", 3)
            add_two_image_row(
                doc,
                male_t if male_t.exists() else None,
                female_t if female_t.exists() else None,
                f"{p['id']}-male.jpg" if male_t.exists() else "(none)",
                f"{p['id']}-female.jpg" if female_t.exists() else "(none)",
            )
            seen_ids.add(p["id"])

    doc.add_page_break()

    # -- API endpoints --------------------------------------------------
    add_heading(doc, "10. API endpoints", 1)

    add_heading(doc, "POST /api/generate — primary generation", 2)
    add_kv_table(doc, [
        ("Method", "POST"),
        ("Body", "multipart/form-data"),
        ("Fields", "userImage (file, ≤30 MB), presetId (string), gender ('male'|'female')"),
        ("Model", "gemini-3-pro-image-preview"),
        ("Output", "{ success, generatedImage: <base64>, mimeType }"),
        ("Failure modes", "400 missing fields · 400 unknown presetId · 500 generation failure"),
        ("Mock mode", "If GEMINI_API_KEY is unset, returns the user photo back with a note."),
        ("Vercel max duration", "60 s"),
    ])

    add_heading(doc, "POST /api/faceswap — optional InstantID identity lock", 2)
    add_paragraph(doc,
        "Not currently wired into the active client flow. Available as a follow-on step when "
        "the Gemini output drifts on identity. Takes the user's webcam photo as sourceImage "
        "and the Gemini scene as targetImage, runs zsxkib/instant-id on Replicate with strong "
        "IdentityNet + IP-Adapter and depth ControlNet to preserve scene framing while "
        "regenerating the person with the user's exact face."
    )
    add_kv_table(doc, [
        ("Method", "POST"),
        ("Body", "multipart/form-data"),
        ("Fields", "sourceImage (face), targetImage (scene), prompt (optional)"),
        ("Model", "zsxkib/instant-id  (SDXL + IdentityNet + IP-Adapter + depth ControlNet)"),
        ("controlnet_conditioning_scale", "0.9 — strong identity match"),
        ("ip_adapter_scale", "0.85 — preserve fine facial features"),
        ("pose_strength", "0.5 — preserve head/body pose from Gemini scene"),
        ("depth_strength", "0.55 — preserve scene framing"),
        ("num_inference_steps", "30"),
        ("Mock mode", "If REPLICATE_API_TOKEN is unset, returns the scene unchanged."),
    ])

    add_heading(doc, "InstantID prompt (default if none provided)", 3)
    add_code(doc, FACESWAP_PROMPT_FALLBACK)

    add_heading(doc, "InstantID negative prompt", 3)
    add_code(doc, FACESWAP_NEGATIVE)

    add_heading(doc, "GET /api/health — config probe", 2)
    add_code(doc, """// api/health.js
module.exports = async function handler(req, res) {
    const geminiConfigured = !!process.env.GEMINI_API_KEY;
    const replicateConfigured = !!process.env.REPLICATE_API_TOKEN;
    res.json({
        status: "ok",
        apiKeyConfigured: geminiConfigured && replicateConfigured,
        geminiConfigured,
        replicateConfigured,
    });
};""")

    doc.add_page_break()

    # -- Frontend deep dive ---------------------------------------------
    add_heading(doc, "11. Frontend — app.js notes", 1)

    add_heading(doc, "State machine", 3)
    add_code(doc, """const state = {
    step: 1,                 // 1 capture · 2 destination · 3 reveal
    stream: null,            // active MediaStream
    capturedBlob: null,      // JPEG blob from canvas
    capturedUrl: null,       // object URL for preview
    selectedPreset: null,    // chosen preset object
    selectedGender: null,    // 'male' | 'female'
    cameraDevices: [],
};""")

    add_heading(doc, "Capture pipeline", 3)
    add_paragraph(doc,
        "Camera ideal resolution is 1920×1080. On capture the canvas is mirrored to match "
        "the live preview (users expect the saved shot to look like what they saw). Output is "
        "JPEG at 0.95 quality. Before generation the blob is recompressed to max 2048 px wide "
        "to keep the multipart body small."
    )

    add_heading(doc, "Gender-restricted destinations", 3)
    add_paragraph(doc,
        "Each preset can carry a 'genders' allowlist (e.g. ['male']). visiblePresets() filters "
        "the grid based on the current gender selection. If the user changes gender after "
        "picking a destination that is no longer available, the selection is cleared."
    )

    add_heading(doc, "Loading hints", 3)
    add_code(doc, """const LOADING_HINTS = [
    'Setting the scene…',
    'Tailoring your outfit…',
    'Matching light and shadows…',
    'Adding the final touches…',
    'Almost there…',
];""")

    add_heading(doc, "WhatsApp share — mobile vs desktop", 3)
    add_paragraph(doc,
        "On mobile, navigator.share with a File array opens the native share sheet "
        "(WhatsApp picks up the image directly). On desktop, wa.me URLs can't carry image "
        "payloads — so the app triggers a download and opens WhatsApp Web with a prefilled "
        "caption text, prompting the user to attach the just-downloaded photo manually."
    )

    doc.add_page_break()

    # -- Deployment -----------------------------------------------------
    add_heading(doc, "12. Deployment", 1)
    add_paragraph(doc, "Production is hosted on Vercel. Local dev uses `npm start` against server.js (legacy Nano-Banana flow).")

    add_heading(doc, "vercel.json", 3)
    add_code(doc, """{
  "functions": {
    "api/generate.js": { "maxDuration": 60 },
    "api/faceswap.js": { "maxDuration": 60 }
  }
}""")

    add_heading(doc, "Environment variables", 3)
    add_kv_table(doc, [
        ("GEMINI_API_KEY", "Google Gemini API key — required for real generation. If absent, /api/generate runs in mock mode and returns the user's photo back."),
        ("REPLICATE_API_TOKEN", "Replicate API token — only needed if you wire /api/faceswap into the flow."),
        ("PORT (local only)", "Defaults to 3000 for server.js dev."),
    ])

    add_heading(doc, "Vercel function notes (multipart quirk)", 3)
    add_paragraph(doc,
        "Vercel's Node runtime pre-buffers the multipart body into req.body as a Buffer. "
        "Multer expects a stream, so when the body is already drained the handler "
        "reconstructs a Readable mimicking IncomingMessage. Without this, multer throws "
        "\"Unexpected end of form\". Both /api/generate and /api/faceswap also export "
        "`config = { api: { bodyParser: false } }` so the runtime hands multer the raw stream."
    )

    doc.add_page_break()

    # -- Appendix: legacy README presets --------------------------------
    add_heading(doc, "Appendix A — Legacy README presets (retired)", 1)
    add_paragraph(doc,
        "The original README and QUICKSTART listed an earlier set of 8 generic 'royal Indian "
        "monument' presets. They are NOT in the current PRESETS map and have been retired in "
        "favour of the Madhya Pradesh / Goa / Kashmir set above. Listed here for reference."
    )
    legacy = [
        ("Royal Palace", "Majestic palace with royal attire"),
        ("Taj Mahal", "Iconic monument backdrop"),
        ("Jaipur Fort", "Rajasthani grandeur"),
        ("Gateway of India", "Mumbai landmark"),
        ("Mysore Palace", "South Indian royalty"),
        ("Red Fort", "Historic Delhi monument"),
        ("Hawa Mahal", "Palace of Winds"),
        ("Amber Fort", "Hilltop fortress"),
    ]
    add_kv_table(doc, legacy)

    add_paragraph(doc, "")  # spacer

    # -- Save -----------------------------------------------------------
    doc.save(OUT)
    print(f"✅ Wrote: {OUT}")
    print(f"   Size: {OUT.stat().st_size / 1024 / 1024:.1f} MB")


if __name__ == "__main__":
    build()
